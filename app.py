import streamlit as st
import os
import json
import re
import tempfile
from datetime import datetime
from typing import TypedDict, Annotated, Optional
from io import BytesIO

st.set_page_config(
    page_title="Inspector BIM",
    page_icon="🏗️",
    layout="wide"
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Mono:wght@400;500&family=IBM+Plex+Sans:wght@300;400;500;600&display=swap');

html, body, [class*="css"] { font-family: 'IBM Plex Sans', sans-serif; }

.cabecalho {
    border-bottom: 2px solid #2d4a2d;
    padding-bottom: 1rem;
    margin-bottom: 1.5rem;
}
.titulo-principal {
    font-size: 1.8rem;
    font-weight: 600;
    color: #e8f5e9;
    letter-spacing: -0.3px;
}
.subtitulo-principal {
    font-size: 0.85rem;
    color: #6b7c6b;
    font-family: 'IBM Plex Mono', monospace;
}
.agente-container {
    background: #0d1a0d;
    border: 1px solid #1e3a1e;
    border-left: 3px solid #4caf50;
    border-radius: 4px;
    padding: 0.9rem 1.1rem;
    margin-bottom: 0.6rem;
    font-family: 'IBM Plex Mono', monospace;
    font-size: 0.82rem;
    color: #a5d6a7;
}
.agente-label {
    font-size: 0.72rem;
    text-transform: uppercase;
    letter-spacing: 0.1em;
    color: #66bb6a;
    font-weight: 500;
    margin-bottom: 0.35rem;
}
.ok-pill {
    display: inline-block;
    background: #0d2b0d;
    color: #66bb6a;
    border: 1px solid #2d5a2d;
    padding: 1px 8px;
    border-radius: 3px;
    font-family: 'IBM Plex Mono', monospace;
    font-size: 0.75rem;
    margin-right: 4px;
}
.nok-pill {
    display: inline-block;
    background: #2b0d0d;
    color: #ef9a9a;
    border: 1px solid #5a2d2d;
    padding: 1px 8px;
    border-radius: 3px;
    font-family: 'IBM Plex Mono', monospace;
    font-size: 0.75rem;
    margin-right: 4px;
}
.metric-bloco {
    background: #0d1a0d;
    border: 1px solid #1e3a1e;
    border-radius: 4px;
    padding: 0.9rem;
    text-align: center;
}
.metric-num {
    font-size: 2rem;
    font-weight: 600;
    color: #81c784;
    font-family: 'IBM Plex Mono', monospace;
    line-height: 1.1;
}
.metric-leg {
    font-size: 0.73rem;
    color: #4a6a4a;
    text-transform: uppercase;
    letter-spacing: 0.07em;
    margin-top: 0.25rem;
}
.relatorio-bloco {
    background: #0d1a0d;
    border: 1px solid #1e3a1e;
    border-radius: 4px;
    padding: 1.2rem 1.4rem;
    font-size: 0.87rem;
    color: #c8e6c9;
    line-height: 1.75;
    font-family: 'IBM Plex Sans', sans-serif;
    white-space: pre-wrap;
}
.secção {
    font-size: 0.72rem;
    text-transform: uppercase;
    letter-spacing: 0.1em;
    color: #4a6a4a;
    font-family: 'IBM Plex Mono', monospace;
    margin-bottom: 0.6rem;
    margin-top: 1.4rem;
}
</style>
""", unsafe_allow_html=True)


# ============================================================
# Cabeçalho
# ============================================================

st.markdown("""
<div class="cabecalho">
    <div class="titulo-principal">Inspector BIM</div>
    <div class="subtitulo-principal">sistema multi-agente / LangGraph + IfcOpenShell / edificações agrícolas</div>
</div>
""", unsafe_allow_html=True)

col_esq, col_dir = st.columns([1, 2], gap="large")

# ============================================================
# Painel esquerdo — configuração
# ============================================================

with col_esq:
    st.markdown('<div class="secção">Configuração</div>', unsafe_allow_html=True)

    api_key = st.text_input(
        "ANTHROPIC_API_KEY",
        type="password",
        placeholder="sk-ant-...",
        help="Necessária para o Agente de Recomendações e o Sintetizador. Os restantes agentes funcionam sem API Key."
    )
    if api_key:
        os.environ["ANTHROPIC_API_KEY"] = api_key

    st.markdown('<div class="secção">Ficheiro IFC</div>', unsafe_allow_html=True)

    ficheiro_ifc = st.file_uploader(
        "Carregar ficheiro .ifc",
        type=["ifc"],
        label_visibility="collapsed"
    )

    if ficheiro_ifc:
        st.markdown(f"""
        <div class="agente-container">
            <div class="agente-label">Ficheiro carregado</div>
            {ficheiro_ifc.name}<br>
            {round(ficheiro_ifc.size / 1024, 1)} KB
        </div>
        """, unsafe_allow_html=True)

    st.markdown("")
    executar = st.button("Executar análise", disabled=(not ficheiro_ifc or not api_key))

    if not api_key:
        st.caption("Introduz a API Key para continuar.")
    if not ficheiro_ifc:
        st.caption("Carrega um ficheiro .ifc para continuar.")


# ============================================================
# Painel direito — execução e resultados
# ============================================================

with col_dir:
    if not executar:
        st.markdown('<div class="secção">Como funciona</div>', unsafe_allow_html=True)
        st.markdown("""
        <div class="agente-container">
            <div class="agente-label">Agente 1 — Extrator IFC</div>
            Abre o ficheiro .ifc com IfcOpenShell e extrai paredes, janelas, portas, lajes e espaços funcionais.
        </div>
        <div class="agente-container">
            <div class="agente-label">Agente 2 — Verificador</div>
            Verifica conformidade com DL 163/2006, Portaria 702/80, DL 173/2005 e DL 347/93.
        </div>
        <div class="agente-container">
            <div class="agente-label">Agente 3 — Quantificador</div>
            Calcula medições: área de janelas, contagem de elementos, tipos de parede.
        </div>
        <div class="agente-container" style="border-left-color: #ff8f00;">
            <div class="agente-label" style="color: #ffa726;">Agente 4 — Recomendações LLM (condicional)</div>
            Invocado pelo Claude apenas se existirem não conformidades.
        </div>
        <div class="agente-container">
            <div class="agente-label">Agente 5 — Sintetizador</div>
            Gera o relatório final via LLM e produz os ficheiros .docx, .xlsx e .json.
        </div>
        """, unsafe_allow_html=True)

    if executar and ficheiro_ifc and api_key:

        # Guardar IFC em ficheiro temporário
        with tempfile.NamedTemporaryFile(delete=False, suffix=".ifc") as tmp:
            tmp.write(ficheiro_ifc.read())
            ifc_path = tmp.name

        try:
            # ── Imports do sistema ─────────────────────────────────────────
            from langgraph.graph import StateGraph, START, END
            from langgraph.graph.message import add_messages
            from langchain_anthropic import ChatAnthropic
            from langchain_core.messages import HumanMessage, AIMessage, SystemMessage
            from docx import Document as DocxDocument
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            import ifcopenshell
            import ifcopenshell.util.element

            llm = ChatAnthropic(model="claude-sonnet-4-20250514", temperature=0, max_tokens=1024)

            # ── Estado ────────────────────────────────────────────────────
            class EstadoInspectorAgricola(TypedDict):
                messages: Annotated[list, add_messages]
                caminho_ifc: str
                elementos_ifc: Optional[dict]
                verificacao: Optional[dict]
                medicoes: Optional[dict]
                recomendacoes_llm: Optional[str]
                saidas: Optional[dict]

            # ── Agente 1 — Extrator IFC ───────────────────────────────────
            def agente_extrator_ifc(estado: EstadoInspectorAgricola) -> dict:
                modelo = ifcopenshell.open(estado["caminho_ifc"])

                projecto_ifc = modelo.by_type("IfcProject")
                nome_projecto = projecto_ifc[0].LongName or projecto_ifc[0].Name if projecto_ifc else "Projecto IFC"
                nome_projecto = re.sub(r'\\X\\([0-9A-F]{2})', lambda m: chr(int(m.group(1), 16)), nome_projecto)

                storeys_raw = modelo.by_type("IfcBuildingStorey")
                pisos = [
                    {"nome": s.Name or "Sem nome", "elevacao": round(s.Elevation or 0.0, 3)}
                    for s in storeys_raw
                ]

                paredes_std = modelo.by_type("IfcWallStandardCase")
                paredes_gen = modelo.by_type("IfcWall")
                todas_paredes = list(paredes_std) + list(paredes_gen)
                tipos_parede = {}
                for p in todas_paredes:
                    tipo_limpo = re.sub(r':\d+$', '', p.Name or "Desconhecido")
                    tipos_parede[tipo_limpo] = tipos_parede.get(tipo_limpo, 0) + 1

                janelas_raw = modelo.by_type("IfcWindow")
                janelas = [
                    {"nome": re.sub(r':\d+$', '', j.Name or "Janela"),
                     "altura_m": round(j.OverallHeight or 0.0, 3),
                     "largura_m": round(j.OverallWidth or 0.0, 3)}
                    for j in janelas_raw
                ]

                portas_raw = modelo.by_type("IfcDoor")
                portas = [
                    {"nome": re.sub(r':\d+$', '', p.Name or "Porta"),
                     "altura_m": round(p.OverallHeight or 0.0, 3),
                     "largura_m": round(p.OverallWidth or 0.0, 3)}
                    for p in portas_raw
                ]

                lajes_raw = modelo.by_type("IfcSlab")
                area_piso_m2 = 0.0
                lajes_piso = []
                for laje in lajes_raw:
                    tipo_laje = getattr(laje, 'PredefinedType', None)
                    if tipo_laje not in ('ROOF', 'BASESLAB'):
                        lajes_piso.append(laje)
                        try:
                            psets = ifcopenshell.util.element.get_psets(laje)
                            area = (
                                psets.get("Qto_SlabBaseQuantities", {}).get("NetArea")
                                or psets.get("Qto_SlabBaseQuantities", {}).get("GrossArea")
                                or psets.get("BaseQuantities", {}).get("NetArea")
                            )
                            if area:
                                area_piso_m2 += area
                        except Exception:
                            pass

                espacos_raw = modelo.by_type("IfcSpace")
                espacos = []
                for e in espacos_raw:
                    nome = e.LongName or e.Name or "Sem nome"
                    nome = re.sub(r'\\X\\([0-9A-F]{2})', lambda m: chr(int(m.group(1), 16)), nome)
                    area = None
                    try:
                        psets = ifcopenshell.util.element.get_psets(e)
                        area = (
                            psets.get("Qto_SpaceBaseQuantities", {}).get("NetFloorArea")
                            or psets.get("BaseQuantities", {}).get("NetFloorArea")
                            or psets.get("Pset_SpaceCommon", {}).get("NetPlannedArea")
                        )
                    except Exception:
                        area = None
                    espacos.append({"numero": e.Name or "-", "nome": nome, "area_m2": area})

                elementos = {
                    "projecto": nome_projecto,
                    "pisos": pisos,
                    "paredes": {"total": len(todas_paredes), "tipos": tipos_parede},
                    "janelas": janelas,
                    "portas": portas,
                    "lajes": {"total": len(lajes_raw)},
                    "area_piso_m2": round(area_piso_m2, 2),
                    "espacos": espacos,
                }

                return {
                    "elementos_ifc": elementos,
                    "messages": [AIMessage(content=
                        f"[Agente Extrator] Modelo: {nome_projecto} | "
                        f"{len(todas_paredes)} paredes | {len(janelas)} janelas | "
                        f"{len(portas)} portas | {len(lajes_raw)} lajes | {len(espacos)} espaços"
                    )]
                }

            # ── Agente 2 — Verificador ────────────────────────────────────
            def agente_verificador(estado: EstadoInspectorAgricola) -> dict:
                dados = estado["elementos_ifc"]
                aprovacoes = []
                nao_conformidades = []

                portas_estreitas = [p for p in dados["portas"] if p["largura_m"] < 0.80]
                if portas_estreitas:
                    nao_conformidades.append(
                        f"Acessibilidade (DL 163/2006) — {len(portas_estreitas)} porta(s) com largura < 0.80m: "
                        + ", ".join(f"{p['nome']} ({p['largura_m']:.2f}m)" for p in portas_estreitas)
                    )
                else:
                    aprovacoes.append(
                        f"Acessibilidade (DL 163/2006) — Todas as {len(dados['portas'])} portas têm largura ≥ 0.80m "
                        f"(mín. encontrado: {min(p['largura_m'] for p in dados['portas']):.2f}m)"
                    )

                area_janelas = sum(j["largura_m"] * j["altura_m"] for j in dados["janelas"])
                area_piso = dados.get("area_piso_m2", 0)
                if area_janelas == 0:
                    nao_conformidades.append("Ventilação (Portaria 702/80) — Nenhuma janela encontrada no modelo")
                elif area_piso > 0:
                    racio = area_janelas / area_piso
                    if racio >= 0.10:
                        aprovacoes.append(
                            f"Ventilação (Portaria 702/80) — Rácio de {racio:.1%} "
                            f"(janelas: {area_janelas:.2f} m² / piso: {area_piso:.2f} m²). Cumpre o mínimo de 10%."
                        )
                    else:
                        nao_conformidades.append(
                            f"Ventilação (Portaria 702/80) — Rácio de {racio:.1%} "
                            f"(janelas: {area_janelas:.2f} m² / piso: {area_piso:.2f} m²). Não cumpre o mínimo de 10%."
                        )
                else:
                    aprovacoes.append(
                        f"Ventilação (Portaria 702/80) — Área de janelas: {area_janelas:.2f} m² "
                        f"({len(dados['janelas'])} janelas). Área de pavimento não exportada — verificar rácio em obra."
                    )

                nomes_espacos = [e["nome"].lower() for e in dados["espacos"]]
                espacos_necessarios = {
                    "oficina ou ferramentaria": ["oficina", "ferramentaria"],
                    "armazém de fitofármacos":  ["fitof", "armazém", "armaz"],
                    "parque de máquinas":       ["parque", "máquinas", "maquinas"],
                    "instalações sanitárias":   ["i.s.", "instalação", "sanitár", "wc"],
                }
                for descricao, palavras_chave in espacos_necessarios.items():
                    encontrado = any(any(kw in nome for kw in palavras_chave) for nome in nomes_espacos)
                    if encontrado:
                        aprovacoes.append(f"Programa funcional — Espaço '{descricao}' identificado no modelo")
                    else:
                        nao_conformidades.append(f"Programa funcional — Espaço '{descricao}' não encontrado no modelo")

                fitof_espacos = [e for e in dados["espacos"] if "fitof" in e["nome"].lower()]
                if fitof_espacos:
                    aprovacoes.append(
                        f"DL 173/2005 — '{fitof_espacos[0]['nome']}' existe como espaço autónomo. "
                        f"Verificar em obra: ventilação directa para o exterior e sinalética obrigatória."
                    )

                tem_is_fem  = any("feminina" in n or "duche f" in n or "cabine f" in n for n in nomes_espacos)
                tem_is_masc = any("masculina" in n or "duche m" in n or "cabine m" in n for n in nomes_espacos)
                if tem_is_fem and tem_is_masc:
                    aprovacoes.append("Higiene (DL 347/93 / Portaria 987/93) — IS Feminina e IS Masculina presentes e separadas")
                elif tem_is_fem or tem_is_masc:
                    nao_conformidades.append("Higiene (DL 347/93 / Portaria 987/93) — Apenas um género de IS identificado; verificar separação")
                else:
                    nao_conformidades.append("Higiene (DL 347/93 / Portaria 987/93) — IS não identificadas no modelo")

                total = len(aprovacoes) + len(nao_conformidades)
                pct = round(len(aprovacoes) / total * 100) if total > 0 else 0
                resumo = f"{len(aprovacoes)} / {total} critérios aprovados ({pct}%)"

                return {
                    "verificacao": {"aprovacoes": aprovacoes, "nao_conformidades": nao_conformidades, "resumo": resumo},
                    "messages": [AIMessage(content=
                        f"[Agente Verificador] {len(aprovacoes)} aprovações | "
                        f"{len(nao_conformidades)} não conformidades | Conformidade: {resumo}"
                    )]
                }

            # ── Agente 3 — Quantificador ──────────────────────────────────
            def agente_quantificador(estado: EstadoInspectorAgricola) -> dict:
                dados = estado["elementos_ifc"]
                area_janelas = sum(j["largura_m"] * j["altura_m"] for j in dados["janelas"])
                tipos_janela = {}
                for j in dados["janelas"]:
                    tipos_janela[j["nome"]] = tipos_janela.get(j["nome"], 0) + 1
                tipos_porta = {}
                for p in dados["portas"]:
                    tipos_porta[p["nome"]] = tipos_porta.get(p["nome"], 0) + 1
                espessuras = sorted(set(
                    int(m) for tipo in dados["paredes"]["tipos"].keys()
                    for m in re.findall(r'(\d+)\s*mm', tipo)
                ))
                medicoes = {
                    "n_janelas": len(dados["janelas"]),
                    "area_janelas_m2": round(area_janelas, 2),
                    "tipos_janela": tipos_janela,
                    "n_portas": len(dados["portas"]),
                    "tipos_porta": tipos_porta,
                    "n_paredes": dados["paredes"]["total"],
                    "tipos_parede": dados["paredes"]["tipos"],
                    "espessuras_mm": espessuras,
                    "n_lajes": dados["lajes"]["total"],
                    "n_espacos": len(dados["espacos"]),
                    "n_pisos": len(dados["pisos"]),
                }
                return {
                    "medicoes": medicoes,
                    "messages": [AIMessage(content=
                        f"[Agente Quantificador] {medicoes['n_janelas']} janelas "
                        f"({medicoes['area_janelas_m2']} m²) | "
                        f"{medicoes['n_portas']} portas | {medicoes['n_paredes']} paredes"
                    )]
                }

            # ── Agente 4 — Recomendações ──────────────────────────────────
            def agente_recomendacoes(estado: EstadoInspectorAgricola) -> dict:
                nao_conformidades = estado["verificacao"]["nao_conformidades"]
                dados = estado["elementos_ifc"]
                system_prompt = """És um engenheiro técnico especializado em edificações agrícolas e industriais em Portugal.
Para cada não conformidade identificada num modelo BIM (IFC), fornece:
1. Acção correctiva concreta e exequível
2. Referência normativa ou regulamentar aplicável (legislação portuguesa)
3. Prazo estimado de resolução
Sê técnico, directo e objectivo. Usa linguagem de relatório técnico em Português de Portugal."""
                nao_conformidades_texto = "\n".join(f"- {p}" for p in nao_conformidades)
                human_prompt = f"""Projecto: {dados['projecto']}
Tipo: Edificação de apoio agrícola (Casa de Máquinas)

Não Conformidades identificadas no modelo BIM:
{nao_conformidades_texto}

Para cada não conformidade, fornece a acção correctiva, referência normativa e prazo."""
                resposta = llm.invoke([
                    SystemMessage(content=system_prompt),
                    HumanMessage(content=human_prompt)
                ])
                return {
                    "recomendacoes_llm": resposta.content,
                    "messages": [AIMessage(content=f"[Agente Recomendações] {resposta.content}")]
                }

            # ── Funções auxiliares do Sintetizador ────────────────────────
            def gerar_docx_bytes(relatorio_txt, dados, verificacao):
                doc = DocxDocument()
                titulo = doc.add_heading(dados["projecto"], level=0)
                titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sub = doc.add_paragraph(f"Relatório de Inspecção BIM  |  {datetime.now().strftime('%d/%m/%Y')}")
                sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
                for linha in relatorio_txt.split("\n"):
                    linha_strip = linha.strip()
                    if not linha_strip:
                        continue
                    if linha_strip.startswith("## "):
                        doc.add_heading(linha_strip[3:], level=1)
                    elif linha_strip.startswith("### "):
                        doc.add_heading(linha_strip[4:], level=2)
                    elif linha_strip.startswith("- ") or linha_strip.startswith("* "):
                        doc.add_paragraph(linha_strip[2:], style="List Bullet")
                    elif linha_strip.startswith("**") and linha_strip.endswith("**"):
                        p = doc.add_paragraph()
                        p.add_run(linha_strip.strip("*")).bold = True
                    else:
                        doc.add_paragraph(linha_strip)
                buf = BytesIO()
                doc.save(buf)
                buf.seek(0)
                return buf.read()

            def gerar_xlsx_bytes(verificacao, medicoes, dados):
                wb = openpyxl.Workbook()
                ws = wb.active
                ws.title = "Conformidade"
                ws.sheet_view.showGridLines = False
                cor_titulo = PatternFill("solid", fgColor="1B5E20")
                cor_ok  = PatternFill("solid", fgColor="E8F5E9")
                cor_nok = PatternFill("solid", fgColor="FFEBEE")
                cor_hdr = PatternFill("solid", fgColor="2E7D32")
                fonte_titulo = Font(bold=True, color="FFFFFF", name="Calibri", size=12)
                fonte_hdr    = Font(bold=True, color="FFFFFF", name="Calibri", size=11)
                al_centro = Alignment(horizontal="center", vertical="center", wrap_text=True)
                al_esq    = Alignment(horizontal="left",   vertical="center", wrap_text=True)
                borda = Border(
                    left=Side(style="thin", color="BDBDBD"), right=Side(style="thin", color="BDBDBD"),
                    top=Side(style="thin", color="BDBDBD"),  bottom=Side(style="thin", color="BDBDBD")
                )
                ws.column_dimensions["A"].width = 12
                ws.column_dimensions["B"].width = 65
                ws.merge_cells("A1:B1")
                c = ws["A1"]
                c.value = f"CHECKLIST DE CONFORMIDADE — {dados['projecto'].upper()}"
                c.fill = cor_titulo; c.font = fonte_titulo; c.alignment = al_centro
                ws.row_dimensions[1].height = 30
                ws.merge_cells("A2:B2")
                c = ws["A2"]
                c.value = f"Inspector BIM  |  {datetime.now().strftime('%d/%m/%Y')}"
                c.fill = PatternFill("solid", fgColor="388E3C")
                c.font = Font(color="FFFFFF", name="Calibri", size=10, italic=True)
                c.alignment = al_centro; ws.row_dimensions[2].height = 18
                for col, hdr in zip(["A", "B"], ["Estado", "Critério / Descrição"]):
                    c = ws[f"{col}3"]
                    c.value = hdr; c.fill = cor_hdr; c.font = fonte_hdr
                    c.alignment = al_centro; c.border = borda
                ws.row_dimensions[3].height = 22
                linha = 4
                for a in verificacao["aprovacoes"]:
                    ws.cell(linha, 1, "OK").fill = cor_ok
                    ws.cell(linha, 1).font = Font(name="Calibri", size=10, color="1B5E20", bold=True)
                    ws.cell(linha, 1).alignment = al_centro; ws.cell(linha, 1).border = borda
                    ws.cell(linha, 2, a).fill = cor_ok
                    ws.cell(linha, 2).font = Font(name="Calibri", size=10, color="1B5E20")
                    ws.cell(linha, 2).alignment = al_esq; ws.cell(linha, 2).border = borda
                    ws.row_dimensions[linha].height = 28; linha += 1
                for n in verificacao["nao_conformidades"]:
                    ws.cell(linha, 1, "NOK").fill = cor_nok
                    ws.cell(linha, 1).font = Font(name="Calibri", size=10, color="B71C1C", bold=True)
                    ws.cell(linha, 1).alignment = al_centro; ws.cell(linha, 1).border = borda
                    ws.cell(linha, 2, n).fill = cor_nok
                    ws.cell(linha, 2).font = Font(name="Calibri", size=10, color="B71C1C")
                    ws.cell(linha, 2).alignment = al_esq; ws.cell(linha, 2).border = borda
                    ws.row_dimensions[linha].height = 32; linha += 1
                ws.merge_cells(f"A{linha}:B{linha}")
                c = ws.cell(linha, 1, f"RESULTADO: {verificacao['resumo']}")
                c.fill = cor_titulo; c.font = fonte_titulo; c.alignment = al_centro
                ws.row_dimensions[linha].height = 26
                ws2 = wb.create_sheet("Medicoes")
                ws2.sheet_view.showGridLines = False
                ws2.column_dimensions["A"].width = 35; ws2.column_dimensions["B"].width = 25
                ws2.merge_cells("A1:B1")
                c = ws2["A1"]; c.value = "MEDICOES"
                c.fill = cor_titulo; c.font = fonte_titulo; c.alignment = al_centro
                ws2.row_dimensions[1].height = 26
                for col, hdr in zip(["A", "B"], ["Elemento", "Quantidade"]):
                    c = ws2[f"{col}2"]; c.value = hdr
                    c.fill = cor_hdr; c.font = fonte_hdr; c.alignment = al_centro; c.border = borda
                ws2.row_dimensions[2].height = 22
                rows = [
                    ("Janelas (unidades)", medicoes["n_janelas"]),
                    ("Janelas (area m2)", medicoes["area_janelas_m2"]),
                    ("Portas (unidades)", medicoes["n_portas"]),
                    ("Paredes (total)", medicoes["n_paredes"]),
                    ("Espessuras de parede (mm)", str(medicoes["espessuras_mm"])),
                    ("Lajes", medicoes["n_lajes"]),
                    ("Espaços funcionais", medicoes["n_espacos"]),
                    ("Pisos", medicoes["n_pisos"]),
                ]
                for i, (elem, val) in enumerate(rows, 3):
                    fill = PatternFill("solid", fgColor="ECEFF1" if i % 2 == 0 else "FFFFFF")
                    ws2.cell(i, 1, elem).fill = fill; ws2.cell(i, 1).alignment = al_esq; ws2.cell(i, 1).border = borda
                    ws2.cell(i, 2, val).fill  = fill; ws2.cell(i, 2).alignment = al_centro; ws2.cell(i, 2).border = borda
                    ws2.row_dimensions[i].height = 20
                buf = BytesIO()
                wb.save(buf); buf.seek(0)
                return buf.read()

            # ── Agente 5 — Sintetizador ───────────────────────────────────
            def agente_sintetizador(estado: EstadoInspectorAgricola) -> dict:
                dados = estado["elementos_ifc"]
                verificacao = estado["verificacao"]
                medicoes = estado["medicoes"]
                mensagens_para_llm = [
                    SystemMessage(content=
                        "És um engenheiro técnico especializado em edificações agrícolas em Portugal. "
                        "Com base nos relatórios dos agentes abaixo, redige um relatório técnico final "
                        "em Português de Portugal. Estrutura o relatório com: "
                        "1. Sumário Executivo, "
                        "2. Dados do Modelo IFC, "
                        "3. Verificação de Conformidade, "
                        "4. Medições, "
                        "5. Recomendações Técnicas (se aplicável), "
                        "6. Conclusão. "
                        "Usa linguagem técnica e formal."
                    )
                ] + estado["messages"]
                resposta = llm.invoke(mensagens_para_llm)
                relatorio = resposta.content
                return {
                    "saidas": {
                        "relatorio_txt": relatorio,
                        "docx_bytes": gerar_docx_bytes(relatorio, dados, verificacao),
                        "xlsx_bytes": gerar_xlsx_bytes(verificacao, medicoes, dados),
                        "json_bytes": json.dumps({
                            "meta": {
                                "sistema": "Inspector BIM — Edificação de Apoio Agricola",
                                "versao": "1.0",
                                "data": datetime.now().isoformat(),
                                "projecto": dados["projecto"],
                            },
                            "elementos_extraidos": dados,
                            "verificacao": verificacao,
                            "medicoes": medicoes,
                            "recomendacoes": estado.get("recomendacoes_llm"),
                        }, ensure_ascii=False, indent=2).encode("utf-8"),
                    },
                    "messages": [AIMessage(content="[Sintetizador] Relatório técnico final gerado.")]
                }

            # ── Roteamento ────────────────────────────────────────────────
            def rotear_apos_quantificador(estado: EstadoInspectorAgricola) -> str:
                if estado["verificacao"]["nao_conformidades"]:
                    return "agente_recomendacoes"
                return "agente_sintetizador"

            # ── Grafo ─────────────────────────────────────────────────────
            builder = StateGraph(EstadoInspectorAgricola)
            builder.add_node("agente_extrator",      agente_extrator_ifc)
            builder.add_node("agente_verificador",   agente_verificador)
            builder.add_node("agente_quantificador", agente_quantificador)
            builder.add_node("agente_recomendacoes", agente_recomendacoes)
            builder.add_node("agente_sintetizador",  agente_sintetizador)
            builder.add_edge(START, "agente_extrator")
            builder.add_edge("agente_extrator",      "agente_verificador")
            builder.add_edge("agente_verificador",   "agente_quantificador")
            builder.add_conditional_edges(
                "agente_quantificador", rotear_apos_quantificador,
                {"agente_recomendacoes": "agente_recomendacoes", "agente_sintetizador": "agente_sintetizador"}
            )
            builder.add_edge("agente_recomendacoes", "agente_sintetizador")
            builder.add_edge("agente_sintetizador", END)
            grafo = builder.compile()

            # ── Execução com feedback em tempo real ───────────────────────
            st.markdown('<div class="secção">Execução</div>', unsafe_allow_html=True)

            placeholder_agentes = st.empty()
            log_agentes = []

            def mostrar_agentes(agentes):
                html = ""
                for a in agentes:
                    cor = "#ff8f00" if "Recomendações" in a["nome"] else "#4caf50"
                    html += f"""
                    <div class="agente-container" style="border-left-color:{cor};">
                        <div class="agente-label" style="color:{cor};">{a["nome"]}</div>
                        {a["msg"]}
                    </div>"""
                placeholder_agentes.markdown(html, unsafe_allow_html=True)

            nomes_agentes = {
                "agente_extrator":      "Agente 1 — Extrator IFC",
                "agente_verificador":   "Agente 2 — Verificador",
                "agente_quantificador": "Agente 3 — Quantificador",
                "agente_recomendacoes": "Agente 4 — Recomendações LLM",
                "agente_sintetizador":  "Agente 5 — Sintetizador",
            }

            msgs_agente = {
                "agente_extrator":      "A ler o ficheiro IFC com IfcOpenShell...",
                "agente_verificador":   "A verificar conformidade normativa...",
                "agente_quantificador": "A calcular medições...",
                "agente_recomendacoes": "A gerar recomendações com o Claude...",
                "agente_sintetizador":  "A gerar relatório final e ficheiros de output...",
            }

            resultado = None
            for evento in grafo.stream(
                {"messages": [], "caminho_ifc": ifc_path,
                 "elementos_ifc": None, "verificacao": None,
                 "medicoes": None, "recomendacoes_llm": None, "saidas": None},
                stream_mode="updates"
            ):
                for nome_no in evento:
                    log_agentes.append({"nome": nomes_agentes.get(nome_no, nome_no), "msg": msgs_agente.get(nome_no, "")})
                    mostrar_agentes(log_agentes)

                if "agente_sintetizador" in evento:
                    resultado = evento["agente_sintetizador"]

            # Reconstruir estado final para aceder aos dados
            estado_final = grafo.invoke(
                {"messages": [], "caminho_ifc": ifc_path,
                 "elementos_ifc": None, "verificacao": None,
                 "medicoes": None, "recomendacoes_llm": None, "saidas": None}
            )

            # ── Resultados ────────────────────────────────────────────────
            st.markdown('<div class="secção">Resultados</div>', unsafe_allow_html=True)

            verificacao = estado_final["verificacao"]
            medicoes    = estado_final["medicoes"]
            dados       = estado_final["elementos_ifc"]
            saidas      = estado_final["saidas"]

            # Métricas
            c1, c2, c3, c4, c5 = st.columns(5)
            with c1:
                st.markdown(f'<div class="metric-bloco"><div class="metric-num">{medicoes["n_paredes"]}</div><div class="metric-leg">Paredes</div></div>', unsafe_allow_html=True)
            with c2:
                st.markdown(f'<div class="metric-bloco"><div class="metric-num">{medicoes["n_janelas"]}</div><div class="metric-leg">Janelas</div></div>', unsafe_allow_html=True)
            with c3:
                st.markdown(f'<div class="metric-bloco"><div class="metric-num">{medicoes["n_portas"]}</div><div class="metric-leg">Portas</div></div>', unsafe_allow_html=True)
            with c4:
                st.markdown(f'<div class="metric-bloco"><div class="metric-num">{medicoes["n_lajes"]}</div><div class="metric-leg">Lajes</div></div>', unsafe_allow_html=True)
            with c5:
                st.markdown(f'<div class="metric-bloco"><div class="metric-num">{medicoes["n_espacos"]}</div><div class="metric-leg">Espaços</div></div>', unsafe_allow_html=True)

            st.markdown("")

            # Conformidade
            st.markdown('<div class="secção">Verificação de Conformidade</div>', unsafe_allow_html=True)
            st.markdown(f"**{verificacao['resumo']}**")
            for a in verificacao["aprovacoes"]:
                st.markdown(f'<span class="ok-pill">OK</span> {a}', unsafe_allow_html=True)
            for n in verificacao["nao_conformidades"]:
                st.markdown(f'<span class="nok-pill">NOK</span> {n}', unsafe_allow_html=True)

            # Relatório
            st.markdown('<div class="secção">Relatório Final</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="relatorio-bloco">{saidas["relatorio_txt"]}</div>', unsafe_allow_html=True)

            # Downloads
            st.markdown('<div class="secção">Descarregar Ficheiros</div>', unsafe_allow_html=True)
            d1, d2, d3 = st.columns(3)
            with d1:
                st.download_button(
                    "Relatório .docx",
                    data=saidas["docx_bytes"],
                    file_name="relatorio_inspector.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            with d2:
                st.download_button(
                    "Checklist .xlsx",
                    data=saidas["xlsx_bytes"],
                    file_name="checklist_conformidade.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
            with d3:
                st.download_button(
                    "Log .json",
                    data=saidas["json_bytes"],
                    file_name="log_inspector.json",
                    mime="application/json"
                )

        except Exception as e:
            st.error(f"Erro durante a execução: {e}")
            raise e
        finally:
            os.unlink(ifc_path)
